"""
Invoice data extraction from uploaded documents (images, PDFs, DOCX).

Supported formats:
  - Images (JPEG, PNG, WEBP, BMP, TIFF) → OCR via pytesseract
  - PDF → text extraction via pdfplumber (falls back to OCR for scanned PDFs)
  - DOCX → text extraction via python-docx

Extraction strategy:
  Uses regex patterns to locate common invoice fields (invoice number, supplier,
  amount, date, description) from the extracted text.  Returns a dict of extracted
  fields that can be used to pre-fill an InvoiceCreate schema.
"""
import io
import logging
import re
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation
from typing import Any

from src.config import settings

logger = logging.getLogger(__name__)


# ── Text extractors ──────────────────────────────────────────────────────────

def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from an image using Tesseract OCR."""
    try:
        from PIL import Image
        import pytesseract
    except ImportError as e:
        raise RuntimeError(
            "pytesseract and Pillow are required for image OCR. "
            "Install with: pip install pytesseract Pillow"
        ) from e

    image = Image.open(io.BytesIO(file_bytes))
    text = pytesseract.image_to_string(image, lang="spa+eng")
    return text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF. Uses pdfplumber for text-based PDFs."""
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError(
            "pdfplumber is required for PDF extraction. "
            "Install with: pip install pdfplumber"
        ) from e

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n".join(text_parts)

    # If no text extracted (scanned PDF), try OCR fallback
    if not full_text.strip():
        logger.info("PDF has no embedded text — attempting OCR fallback")
        try:
            full_text = _ocr_pdf_pages(file_bytes)
        except Exception:
            logger.warning("OCR fallback for PDF failed", exc_info=True)

    return full_text


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    """OCR fallback for scanned PDFs — convert pages to images and run tesseract."""
    try:
        from PIL import Image
        import pytesseract
        import pdfplumber
    except ImportError:
        return ""

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            img = page.to_image(resolution=300).original
            page_text = pytesseract.image_to_string(img, lang="spa+eng")
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install python-docx"
        ) from e

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


# ── Field extraction via regex ───────────────────────────────────────────────

# Invoice number patterns (FAC-001, Factura N° 12345, Invoice #ABC-123, No. 001, etc.)
_INVOICE_NUMBER_PATTERNS = [
    r"(?:factura|invoice|fact|fac|nota|recibo|comprobante)\s*(?:n[°ºo.]?|#|num|número|numero)?\s*[:\-]?\s*([A-Za-z0-9][\w\-/]{1,30})",
    r"(?:n[°ºo.]|#|no\.?)\s*[:\-]?\s*([A-Za-z0-9][\w\-/]{2,30})",
    r"\b(FAC[\-/]?\d{2,10})\b",
    r"\b(INV[\-/]?\d{2,10})\b",
    r"\b(FACT[\-/]?\d{2,10})\b",
]

# Supplier / company name patterns
_SUPPLIER_PATTERNS = [
    r"(?:proveedor|supplier|empresa|razón social|razon social|vendedor|emitido por|de:)\s*[:\-]?\s*(.+)",
    r"(?:NIT|RUC|RUT|RFC)\s*[:\-]?\s*[\d.\-]+\s*[;\n]?\s*(.+)",
]

# Amount / total patterns (handles $1.500.000,50 or $1,500,000.50 or 1500000.50)
_AMOUNT_PATTERNS = [
    r"(?:total\s*(?:a pagar|factura|general|neto)?|monto\s*total|gran total|valor total|importe total|amount|total due)\s*[:\-]?\s*\$?\s*([\d.,]+(?:\.\d{1,2})?)",
    r"(?:subtotal|sub[\- ]total)\s*[:\-]?\s*\$?\s*([\d.,]+(?:\.\d{1,2})?)",
    r"(?:total)\s*[:\-]?\s*(?:COP|USD|EUR)?\s*\$?\s*([\d.,]+(?:\.\d{1,2})?)",
]

# Date patterns (DD/MM/YYYY, DD-MM-YYYY, YYYY-MM-DD, etc.)
_DATE_PATTERNS = [
    r"(?:fecha\s*(?:de\s*)?(?:vencimiento|pago|límite|vence)|due date|payment date|vencimiento)\s*[:\-]?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
    r"(?:fecha\s*(?:de\s*)?(?:emisión|emision|factura|facturación|expedición)|date|invoice date)\s*[:\-]?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
    r"\b(\d{4}[/\-]\d{1,2}[/\-]\d{1,2})\b",
    r"\b(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})\b",
]

# Description patterns
_DESCRIPTION_PATTERNS = [
    r"(?:descripción|descripcion|concepto|detalle|description|observaciones|nota)\s*[:\-]?\s*(.+)",
]


def _search_patterns(text: str, patterns: list[str]) -> str | None:
    """Try each regex pattern and return the first match (group 1)."""
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return None


def _parse_amount(raw: str) -> Decimal | None:
    """
    Parse a raw amount string into a Decimal.
    Handles formats like:
      1.500.000,50  →  1500000.50  (European/LatAm)
      1,500,000.50  →  1500000.50  (US)
      1500000.50    →  1500000.50
      1500000       →  1500000
    """
    if not raw:
        return None

    raw = raw.strip().replace(" ", "")

    # Detect format: if last separator is comma and ≤2 digits follow → comma = decimal
    if re.match(r"^[\d.]+,\d{1,2}$", raw):
        # European/LatAm: 1.500.000,50
        raw = raw.replace(".", "").replace(",", ".")
    elif re.match(r"^[\d,]+\.\d{1,2}$", raw):
        # US: 1,500,000.50
        raw = raw.replace(",", "")
    else:
        # Remove thousand separators (could be dots or commas)
        raw = raw.replace(",", "").replace(".", "")
        # If we stripped everything to digits only, it's a whole number
        if not raw:
            return None

    try:
        amount = Decimal(raw)
        return amount if amount > 0 else None
    except (InvalidOperation, ValueError):
        return None


def _parse_date(raw: str) -> datetime | None:
    """Try to parse a date string into a datetime object."""
    if not raw:
        return None

    formats = [
        "%d/%m/%Y", "%d-%m-%Y",
        "%Y-%m-%d", "%Y/%m/%d",
        "%d/%m/%y", "%d-%m-%y",
        "%m/%d/%Y", "%m-%d-%Y",
    ]
    for fmt in formats:
        try:
            dt = datetime.strptime(raw.strip(), fmt)
            return dt.replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    return None


def extract_invoice_data(text: str) -> dict[str, Any]:
    """
    Extract structured invoice fields from raw text.

    Returns a dict with keys that may include:
      - invoice_number: str | None
      - supplier: str | None
      - amount: float | None
      - description: str | None
      - due_date: str (ISO format) | None
    Only non-None values are included.
    """
    result: dict[str, Any] = {}

    # Invoice number
    inv_num = _search_patterns(text, _INVOICE_NUMBER_PATTERNS)
    if inv_num:
        result["invoice_number"] = inv_num

    # Supplier
    supplier = _search_patterns(text, _SUPPLIER_PATTERNS)
    if supplier:
        # Clean up: take first line, trim excessive length
        supplier = supplier.split("\n")[0].strip()
        if len(supplier) > 200:
            supplier = supplier[:200]
        result["supplier"] = supplier

    # Amount — try total first, then subtotal
    amount_raw = _search_patterns(text, _AMOUNT_PATTERNS)
    if amount_raw:
        amount = _parse_amount(amount_raw)
        if amount:
            result["amount"] = float(amount)

    # Due date
    due_date_raw = _search_patterns(text, _DATE_PATTERNS[:2])  # Only vencimiento patterns
    if due_date_raw:
        dt = _parse_date(due_date_raw)
        if dt:
            result["due_date"] = dt.isoformat()

    # Description
    desc = _search_patterns(text, _DESCRIPTION_PATTERNS)
    if desc:
        desc = desc.split("\n")[0].strip()
        if len(desc) > 500:
            desc = desc[:500]
        result["description"] = desc

    return result


def merge_extractions(regex_data: dict[str, Any], gemini_data: dict[str, Any]) -> dict[str, Any]:
    """Prioriza valores no vacíos de Gemini sobre el extractor por regex."""
    merged = dict(regex_data)
    for key, value in gemini_data.items():
        if value is None:
            continue
        if isinstance(value, str) and not value.strip():
            continue
        if key == "amount" and isinstance(value, (int, float)) and value <= 0:
            continue
        merged[key] = value
    return merged


# ── Main entry point ─────────────────────────────────────────────────────────

SUPPORTED_IMAGE_TYPES = {
    "image/jpeg", "image/png", "image/webp", "image/bmp", "image/tiff",
}
SUPPORTED_PDF_TYPES = {"application/pdf"}
SUPPORTED_DOC_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
}

ALL_SUPPORTED_TYPES = SUPPORTED_IMAGE_TYPES | SUPPORTED_PDF_TYPES | SUPPORTED_DOC_TYPES


def extract_from_file(file_bytes: bytes, content_type: str, filename: str) -> dict[str, Any]:
    """
    Main entry point: extract invoice data from an uploaded file.

    Args:
        file_bytes: raw file content
        content_type: MIME type of the file
        filename: original filename (used as fallback for type detection)

    Returns:
        dict with extracted invoice fields + 'raw_text' key with the full extracted text.

    Raises:
        ValueError: if the file type is not supported
        RuntimeError: if a required library is not installed
    """
    # Fallback type detection by extension
    if not content_type or content_type == "application/octet-stream":
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        type_map = {
            "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
            "webp": "image/webp", "bmp": "image/bmp", "tiff": "image/tiff",
            "tif": "image/tiff", "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        content_type = type_map.get(ext, content_type)

    if content_type in SUPPORTED_IMAGE_TYPES:
        raw_text = extract_text_from_image(file_bytes)
    elif content_type in SUPPORTED_PDF_TYPES:
        raw_text = extract_text_from_pdf(file_bytes)
    elif content_type in SUPPORTED_DOC_TYPES:
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Tipo de archivo no soportado: {content_type}. "
            f"Formatos aceptados: imágenes (JPEG, PNG), PDF, DOCX."
        )

    logger.info("Extracted %d characters from %s (%s)", len(raw_text), filename, content_type)

    extracted = extract_invoice_data(raw_text)
    extraction_method = "regex"

    if (settings.gemini_api_key or "").strip():
        try:
            from src.extraction_gemini import extract_with_gemini

            gemini_data = extract_with_gemini(file_bytes, content_type, filename, raw_text)
            if gemini_data:
                extracted = merge_extractions(extracted, gemini_data)
                extraction_method = "gemini"
        except Exception as e:
            logger.warning("Gemini extraction failed: %s", e, exc_info=True)

    extracted["raw_text"] = raw_text
    extracted["extraction_method"] = extraction_method

    return extracted
